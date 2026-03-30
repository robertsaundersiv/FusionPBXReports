"""
Generate docs/CALCULATIONS.docx — documents every metric calculation and every
toggle available in the FusionPBX Analytics Dashboard.

Run from the workspace root:
    python scripts/generate_docs.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    h.paragraph_format.space_after = Pt(4)
    return h


def add_para(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_table_row(table, label, value, shaded=False):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value
    if shaded:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F2F2F2")
            tcPr.append(shd)


def add_definition_table(doc, rows):
    """rows: list of (label, value) tuples"""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Detail"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for i, (label, value) in enumerate(rows):
        add_table_row(table, label, value, shaded=(i % 2 == 0))
    doc.add_paragraph()  # spacer


# ---------------------------------------------------------------------------
# Document assembly
# ---------------------------------------------------------------------------

def build_document():
    doc = Document()

    # ── Title ──────────────────────────────────────────────────────────────
    title = doc.add_heading("FusionPBX Analytics Dashboard", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Calculations Reference & Toggle Guide")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)

    doc.add_paragraph()

    add_para(
        doc,
        "This document explains how every metric in the dashboard is computed "
        "and describes each filter toggle available to operators. It is intended "
        "as a reference for analysts, supervisors, and developers who need to "
        "understand or audit the reported numbers.",
        size=10,
    )

    doc.add_page_break()

    # ======================================================================
    # SECTION 1 — CORE CONCEPTS
    # ======================================================================
    set_heading(doc, "1. Core Concepts", 1)

    # 1.1 Data source
    set_heading(doc, "1.1  Data Source", 2)
    add_para(
        doc,
        "All metrics are derived from FusionPBX CDR (Call Detail Record) records "
        "stored in the local PostgreSQL database. The ETL worker fetches new CDRs "
        "from the FusionPBX REST API every 5 minutes and upserts them by "
        "xml_cdr_uuid (idempotent). A 15-minute safety window is re-fetched on "
        "every run to catch late-arriving records.",
    )

    # 1.2 Queue entry deduplication
    set_heading(doc, "1.2  Queue Entry Deduplication", 2)
    add_para(
        doc,
        "A single caller interaction with a queue can generate multiple CDR legs "
        "(e.g. parallel ring, transfer, race condition). To avoid double counting, "
        "all queue metrics group legs by the composite key:",
    )
    add_bullet(doc, "(caller_id_number, cc_queue_joined_epoch)")
    add_para(
        doc,
        "Every group is counted as one queue entry. The classification functions "
        "(answered, abandoned, ASA, AHT) operate on the entire group of legs and "
        "return a single verdict.",
    )

    # 1.3 Timestamps
    set_heading(doc, "1.3  Epoch Timestamps", 2)
    add_para(
        doc,
        "FusionPBX stores all times as Unix epoch integers (seconds since "
        "1970-01-01 00:00:00 UTC). The dashboard stores and compares times in "
        "this native format to avoid time-zone ambiguity during arithmetic. "
        "Display conversions use the operator-selected timezone (default: "
        "America/Phoenix).",
    )

    # 1.4 Date window
    set_heading(doc, "1.4  Date Range Windows", 2)
    add_para(doc, "Preset date ranges resolve to local calendar days in the selected timezone:")
    rows = [
        ("Today", "Current calendar day: 00:00:00 → 23:59:59 (local)"),
        ("Yesterday", "Previous calendar day: 00:00:00 → 23:59:59 (local)"),
        ("Last 7 days", "Today minus 6 days (00:00:00) → today (23:59:59). Covers 7 full calendar days including today."),
        ("Last 30 days", "Today minus 29 days (00:00:00) → today (23:59:59). Covers 30 full calendar days."),
        ("Custom", "Operator-entered start and end dates. Start snaps to 00:00:00; End snaps to 23:59:59."),
    ]
    add_definition_table(doc, rows)

    # ======================================================================
    # SECTION 2 — QUEUE METRICS
    # ======================================================================
    set_heading(doc, "2. Queue Metrics", 1)

    add_para(
        doc,
        "Queue metrics appear on the Queue Performance and Queue Performance "
        "Report pages. All queue metrics are inbound-only; the direction filter "
        "is fixed to 'inbound' and is not exposed to the operator.",
    )

    # 2.1 Offered
    set_heading(doc, "2.1  Offered (Total Queue Calls)", 2)
    add_definition_table(doc, [
        ("Definition", "Number of unique caller interactions that entered a queue."),
        ("Filter", "cc_queue_joined_epoch IS NOT NULL AND direction = 'inbound'"),
        ("Unit", "Count of unique (caller_id_number, cc_queue_joined_epoch) groups"),
        ("Notes", "Any call that reached the queue is counted, regardless of outcome."),
    ])

    # 2.2 Answered
    set_heading(doc, "2.2  Answered", 2)
    add_para(
        doc,
        "A queue entry is classified as answered when any leg in the group "
        "satisfies the answered condition. The exact condition depends on the "
        "Strict Queue Answered Mode toggle (see Section 4.3).",
    )

    set_heading(doc, "Standard mode (toggle OFF — default)", 3)
    add_para(doc, "A queue entry is answered when ANY of the following is true for any leg:")
    add_bullet(doc, "cc_queue_answered_epoch IS NOT NULL  — FusionPBX explicitly marked the queue interaction answered.")
    add_bullet(
        doc,
        "answer_epoch IS NOT NULL  AND  hangup_cause = 'NORMAL_CLEARING'  — "
        "The leg connected (answer_epoch set) and ended cleanly. This catches "
        "answered interactions where cc_queue_answered_epoch was not written by "
        "FusionPBX.",
    )

    set_heading(doc, "Strict mode (toggle ON)", 3)
    add_para(doc, "A queue entry is answered when ANY leg has:")
    add_bullet(doc, "cc_queue_answered_epoch IS NOT NULL  — only the authoritative FusionPBX queue-answered timestamp is accepted.")
    add_para(
        doc,
        "The answer_epoch + NORMAL_CLEARING fallback is disabled in strict mode. "
        "Strict mode produces counts that most closely match the FusionPBX built-in "
        "queue statistics report.",
    )

    # 2.3 Abandoned
    set_heading(doc, "2.3  Abandoned", 2)
    add_para(
        doc,
        "A queue entry is classified as abandoned when it is NOT answered (using "
        "whichever answered logic is active) AND at least one leg in the group "
        "matches ALL of the following criteria:",
    )
    add_bullet(doc, "billsec = 0  (no actual talk time on this leg)")
    add_bullet(doc, "last_app ≠ 'deflect'  (not routed to voicemail or an external destination)")
    add_bullet(doc, "hangup_cause ≠ 'BLIND_TRANSFER'  (not transferred away)")
    add_para(doc, "AND at least one of:")
    add_bullet(doc, "hangup_cause IN ('ORIGINATOR_CANCEL', 'NO_ANSWER', 'USER_BUSY')", level=1)
    add_bullet(doc, "cc_cause = 'TIMEOUT'", level=1)
    add_bullet(doc, "call_disposition = 'missed'", level=1)
    add_para(
        doc,
        "Entries that are neither answered nor abandoned (e.g. deflected to "
        "voicemail, blind-transferred) are counted under 'Other' in the call "
        "outcomes breakdown chart.",
    )

    # 2.4 Answer Rate
    set_heading(doc, "2.4  Answer Rate", 2)
    add_definition_table(doc, [
        ("Formula", "answered ÷ offered × 100"),
        ("Unit", "%"),
        ("Example", "85 answered ÷ 100 offered = 85.0 %"),
        ("Green threshold", "≥ 80 %"),
        ("Yellow threshold", "≥ 70 %"),
    ])

    # 2.5 Abandon Rate
    set_heading(doc, "2.5  Abandon Rate", 2)
    add_definition_table(doc, [
        ("Formula", "abandoned ÷ offered × 100"),
        ("Unit", "%"),
        ("Example", "10 abandoned ÷ 100 offered = 10.0 %"),
        ("Green threshold", "≤ 10 %"),
        ("Yellow threshold", "≤ 20 %"),
    ])

    # 2.6 ASA
    set_heading(doc, "2.6  Average Speed of Answer (ASA)", 2)
    add_para(doc, "Measures how quickly calls are picked up after entering the queue.")
    add_para(doc, "For each answered queue entry the wait time is computed as:")
    add_bullet(
        doc,
        "Standard mode: cc_queue_answered_epoch – cc_queue_joined_epoch  (preferred). "
        "If cc_queue_answered_epoch is NULL but the entry qualifies as answered via the "
        "fallback, answer_epoch – cc_queue_joined_epoch is used instead.",
    )
    add_bullet(
        doc,
        "Strict mode: cc_queue_answered_epoch – cc_queue_joined_epoch only. "
        "Legs without cc_queue_answered_epoch contribute no wait time.",
    )
    add_para(
        doc,
        "When multiple legs exist for one entry, the shortest non-negative wait "
        "time across all legs is used (avoids inflating ASA with race-condition legs).",
    )
    add_definition_table(doc, [
        ("Formula", "AVG(min_wait_per_entry) over answered entries"),
        ("Unit", "Seconds"),
        ("P90", "90th percentile of per-entry wait times (worst 10% of calls)"),
        ("Green threshold", "≤ 30 s"),
        ("Yellow threshold", "≤ 60 s"),
    ])

    # 2.7 Service Level
    set_heading(doc, "2.7  Service Level %", 2)
    add_definition_table(doc, [
        ("Definition", "Percentage of answered calls where the caller waited ≤ threshold seconds."),
        ("Default threshold", "30 seconds (configurable per queue in Admin Settings)"),
        ("Formula", "COUNT(answered entries where wait ≤ threshold) ÷ COUNT(answered entries) × 100"),
        ("Denominator", "Only answered entries are included; unanswered entries do not affect service level."),
        ("Green threshold", "≥ 80 %"),
        ("Yellow threshold", "≥ 70 %"),
    ])

    # 2.8 AHT
    set_heading(doc, "2.8  Average Handle Time (AHT)", 2)
    add_para(doc, "Handle time = talk time + hold time for a single answered queue entry.")
    add_para(doc, "Computation per leg:")
    add_bullet(doc, "duration = billsec + hold_accum_seconds")
    add_para(
        doc,
        "When multiple legs exist for the same entry, the maximum leg duration is "
        "used (avoids double-counting talk time across transferred legs).",
    )
    add_definition_table(doc, [
        ("Formula", "AVG(max_duration_per_answered_entry)"),
        ("Unit", "Seconds"),
        ("P90", "90th percentile of per-entry AHT values (worst 10% of calls)"),
        ("Answered check", "Uses the same answered logic as the Answered metric (strict or standard mode)."),
        ("Green threshold", "≤ 300 s"),
        ("Yellow threshold", "≤ 600 s"),
    ])

    # 2.9 MOS
    set_heading(doc, "2.9  Mean Opinion Score (MOS)", 2)
    add_definition_table(doc, [
        ("Source field", "rtp_audio_in_mos — populated by FusionPBX from RTP statistics"),
        ("Formula", "AVG(rtp_audio_in_mos) WHERE rtp_audio_in_mos > 0"),
        ("Unit", "Score 0–5 (ITU-T P.800 scale; 4.0+ is good quality)"),
        ("P10", "10th percentile — represents the worst quality experienced by 10% of callers"),
        ("Green threshold", "≥ 4.0"),
        ("Yellow threshold", "≥ 3.8"),
    ])

    # 2.10 Callbacks
    set_heading(doc, "2.10  Callbacks", 2)
    add_definition_table(doc, [
        ("Callbacks Offered", "COUNT(*) WHERE cc_agent_type = 'callback'"),
        ("Callbacks Answered", "COUNT(*) WHERE cc_agent_type = 'callback' AND cc_queue_answered_epoch IS NOT NULL"),
        ("Completion Rate", "callbacks_answered ÷ callbacks_offered × 100"),
    ])

    # 2.11 Hourly Buckets
    set_heading(doc, "2.11  Hourly Time-Series Buckets", 2)
    add_para(
        doc,
        "Queue performance charts break the date range into 1-hour buckets. For "
        "each bucket the following values are computed using the same entry-based "
        "deduplication logic as the summary metrics:",
    )
    add_bullet(doc, "offered — unique entries with cc_queue_joined_epoch falling in the hour")
    add_bullet(doc, "answered — answered entries in the hour")
    add_bullet(doc, "abandoned — abandoned entries in the hour")
    add_bullet(doc, "answer_rate, abandon_rate — derived from the above")
    add_bullet(doc, "avg_asa, avg_aht — entry-level averages for the hour")
    add_bullet(doc, "service_level — % of answered entries in the hour with wait ≤ 30 s")

    # 2.12 Heatmaps
    set_heading(doc, "2.12  Heatmap Data (Hour × Day of Week)", 2)
    add_para(
        doc,
        "The offered, abandon-rate, and ASA heatmaps on the Queue Performance page "
        "are derived from raw CDR legs (not deduplicated entries) bucketed by the "
        "local hour and weekday of each record's start_epoch. This reflects raw "
        "distribution patterns across the time grid.",
    )

    # ======================================================================
    # SECTION 3 — AGENT METRICS
    # ======================================================================
    set_heading(doc, "3. Agent Metrics", 1)

    add_para(
        doc,
        "Agent metrics appear on the Agent Performance & Coaching page (leaderboard "
        "and individual agent drilldowns).",
    )

    # 3.1 Agent resolution
    set_heading(doc, "3.1  Agent Identity Resolution", 2)
    add_para(
        doc,
        "A CDR record is attributed to an agent through a multi-step resolution "
        "chain evaluated in order:",
    )
    add_bullet(doc, "1. cc_agent_uuid — direct match to known agent UUID in the Agent table", level=0)
    add_bullet(doc, "2. cc_agent — agent name/identifier in cc_agent field", level=0)
    add_bullet(doc, "3. extension_uuid — extension UUID matched against Extension table", level=0)
    add_bullet(doc, "4. caller_id_number — numeric caller-ID matched against known agent extension numbers", level=0)
    add_bullet(doc, "5. caller_id_name — caller label matched by exact name, embedded extension number, or fuzzy word overlap", level=0)
    add_bullet(doc, "6. Raw fallback — normalized cc_agent / cc_agent_uuid value used as-is", level=0)
    add_para(
        doc,
        "All aliases (UUID, extension, contact, name) for each agent are resolved "
        "to a single canonical agent ID, so records with different identifier "
        "styles are correctly merged.",
    )

    # 3.2 Handled calls
    set_heading(doc, "3.2  Handled Calls", 2)
    add_para(
        doc,
        "A call is counted as handled by an agent when the CDR record for that "
        "agent leg passes the is_handled() check:",
    )
    add_bullet(doc, "status = 'answered'  OR  billsec > 0")
    add_para(
        doc,
        "Multi-leg interactions are further deduplicated by call key "
        "(xml_cdr_uuid / bridge_uuid / caller+time) so each physical call is "
        "counted at most once per agent—the leg with the highest rank (most "
        "complete CDR fields) is kept.",
    )

    # 3.3 Missed calls
    set_heading(doc, "3.3  Missed Calls", 2)
    add_para(doc, "A record contributes to missed calls when:")
    add_bullet(doc, "missed_call = True  OR  cc_cancel_reason = 'TIMEOUT'  OR  call_disposition = 'missed'  OR  hangup_cause = 'NO_ANSWER'")
    add_bullet(doc, "direction is 'inbound' or 'local' (outbound records are never counted as missed).")
    add_bullet(doc, "When 'Include outbound' is ON, records attributed only via caller_id_number / caller_id_name are excluded from missed counts to avoid false positives.")
    add_para(doc, "Missed calls are visible only to Super Admin users.")

    # 3.4 Deflects
    set_heading(doc, "3.4  Deflected Calls", 2)
    add_para(
        doc,
        "A call is classified as a deflect when it was routed away without a live "
        "agent answering it (e.g. voicemail, auto-attendant IVR, external transfer):",
    )
    add_bullet(doc, "last_app = 'deflect'  OR")
    add_bullet(doc, "hangup_cause = 'BLIND_TRANSFER'  OR")
    add_bullet(doc, "cc_cause = 'CANCEL'")
    add_para(
        doc,
        "The 'Exclude voicemail/external deflects' toggle (see Section 4.2) removes "
        "deflected calls from agent handling counts.",
    )

    # 3.5 AHT (agent)
    set_heading(doc, "3.5  Agent AHT", 2)
    add_definition_table(doc, [
        ("Formula", "AVG(billsec + hold_accum_seconds) across handled (answered) calls for the agent"),
        ("Unit", "Seconds"),
    ])

    # 3.6 Outbound
    set_heading(doc, "3.6  Outbound Call Attribution", 2)
    add_para(
        doc,
        "When 'Include outbound' is ON, the leaderboard also counts outbound CDR "
        "records. Outbound records often lack cc_agent_uuid / cc_agent / "
        "extension_uuid; they are matched to agents via caller_id_number (the "
        "agent's extension number). This broadens attribution but may include "
        "personal calls placed from an agent's extension.",
    )

    # ======================================================================
    # SECTION 4 — FILTER TOGGLES
    # ======================================================================
    set_heading(doc, "4. Filter Toggles", 1)

    add_para(
        doc,
        "The filter bar exposes several toggles that change which data is included "
        "in calculations. The toggles available depend on the page:",
    )

    # 4.1 Include outbound
    set_heading(doc, "4.1  Include Outbound  (Agent Performance page)", 2)
    add_definition_table(doc, [
        ("Default", "OFF"),
        ("Pages", "Agent Performance & Coaching — leaderboard and drilldowns"),
        ("Effect when OFF",
         "Only inbound and local CDR records are included. Agents are identified "
         "via their call-center identifiers (cc_agent_uuid, cc_agent, extension_uuid)."),
        ("Effect when ON",
         "Outbound CDR records are also fetched and attributed to agents by "
         "matching caller_id_number against known agent extension numbers. "
         "Outbound handled calls contribute to the agent's handled count and AHT."),
        ("Badge text",
         "When ON, a green badge shows the count of outbound handled calls "
         "included using caller-number attribution."),
        ("Missed calls",
         "Missed call counts are intentionally NOT modified by this toggle to "
         "avoid false positives from outbound-attributed records."),
    ])

    # 4.2 Exclude voicemail/external deflects
    set_heading(doc, "4.2  Exclude Voicemail / External Deflects  (Agent Performance page)", 2)
    add_definition_table(doc, [
        ("Default", "ON (deflects excluded by default)"),
        ("Pages", "Agent Performance & Coaching"),
        ("Effect when ON",
         "Calls classified as deflects (last_app = 'deflect', hangup_cause = "
         "'BLIND_TRANSFER', cc_cause = 'CANCEL') are removed from the agent's "
         "handled call count and AHT computation."),
        ("Effect when OFF",
         "Deflected calls are counted as handled. This inflates handled counts "
         "and can reduce apparent AHT because deflect legs tend to be short."),
        ("Typical use",
         "Keep ON to focus on true agent-assisted calls. Turn OFF to audit the "
         "raw volume of calls touching an agent's endpoint including voicemail."),
    ])

    # 4.3 Strict queue answered mode
    set_heading(doc, "4.3  Strict Queue Answered Mode  (Queue pages)", 2)
    add_definition_table(doc, [
        ("Default", "OFF"),
        ("Pages", "Queue Performance, Queue Performance Report"),
        ("Effect when OFF — Standard mode",
         "A queue entry is answered if cc_queue_answered_epoch IS NOT NULL  OR  "
         "(answer_epoch IS NOT NULL AND hangup_cause = 'NORMAL_CLEARING'). "
         "This 'loose' mode catches interactions that FusionPBX did not mark with "
         "cc_queue_answered_epoch but where a live connection clearly occurred."),
        ("Effect when ON — Strict mode",
         "Only cc_queue_answered_epoch IS NOT NULL is accepted. The answer_epoch "
         "fallback is disabled. This mode closely matches the FusionPBX built-in "
         "queue statistics report and is recommended for SLA/compliance reporting."),
        ("Impact on answered count",
         "Strict mode typically produces lower answered counts than standard mode "
         "because some multi-leg entries with NORMAL_CLEARING but without "
         "cc_queue_answered_epoch are no longer counted. "
         "Based on test data: standard mode absolute answered delta vs FusionPBX = ~1044; "
         "strict mode delta = ~189."),
        ("Impact on ASA",
         "In strict mode ASA is computed only from cc_queue_answered_epoch timestamps. "
         "In standard mode, answer_epoch is used as a fallback when "
         "cc_queue_answered_epoch is absent, which may slightly alter the average."),
        ("Impact on AHT",
         "AHT uses the same answered classification, so entries counted only in "
         "standard mode also contribute their duration to AHT in standard mode. "
         "Their exclusion in strict mode changes the average."),
        ("Impact on abandoned",
         "Because 'abandoned = not answered AND caller-hung-up', switching modes "
         "also changes which entries are eligible to be classified as abandoned. "
         "Standard mode may classify some answered-via-fallback entries differently."),
        ("Recommendation",
         "Use Strict mode when comparing to FusionPBX native reports. Use Standard "
         "mode to get a fuller picture of calls that were genuinely handled even "
         "where cc_queue_answered_epoch was not populated by FusionPBX."),
    ])

    # 4.4 Direction
    set_heading(doc, "4.4  Direction Filter  (Executive Overview)", 2)
    add_definition_table(doc, [
        ("Default", "All directions"),
        ("Pages", "Executive Overview"),
        ("Options", "All Directions / Inbound / Outbound / Local"),
        ("Effect", "Filters the base CDR query to include only records with the matching direction field."),
        ("Queue pages", "Queue Performance and Queue Performance Report pages are ALWAYS inbound-only. The direction dropdown is hidden on those pages and the backend enforces direction = 'inbound' regardless."),
    ])

    # 4.5 Date range timezone
    set_heading(doc, "4.5  UTC Timezone Toggle", 2)
    add_definition_table(doc, [
        ("Default", "Browser local timezone (reported by Intl.DateTimeFormat)"),
        ("Pages", "All pages"),
        ("Effect when OFF",
         "Date preset midnight/end-of-day boundaries are computed in the browser's "
         "local timezone and sent to the API as timezone-aware datetimes. The API "
         "also labels trend chart dates using the local timezone."),
        ("Effect when ON",
         "All boundaries are computed in UTC. Useful when the server and reporting "
         "audience are in different timezones or when comparing to UTC-based "
         "external reports."),
    ])

    # ======================================================================
    # SECTION 5 — EXECUTIVE OVERVIEW METRICS
    # ======================================================================
    set_heading(doc, "5. Executive Overview Metrics", 1)

    add_para(
        doc,
        "The Executive Overview page aggregates queue and agent data across the "
        "selected date range and queue filters.",
    )

    rows = [
        ("Total Inbound Calls",
         "COUNT(*) WHERE direction = 'inbound'. Counts every inbound CDR leg, "
         "not deduplicated by queue entry."),
        ("Total Queue Offered",
         "Count of unique (caller_id_number, cc_queue_joined_epoch) groups WHERE "
         "cc_queue_joined_epoch IS NOT NULL."),
        ("Answered",
         "Count of unique queue entries classified as answered (standard mode — "
         "cc_queue_answered_epoch IS NOT NULL OR answer_epoch + NORMAL_CLEARING)."),
        ("Abandoned",
         "Count of unique queue entries classified as abandoned (not answered AND "
         "caller hung up, excluding deflects and blind transfers)."),
        ("Answer Rate", "answered ÷ offered × 100"),
        ("Abandon Rate", "abandoned ÷ offered × 100"),
        ("Average ASA",
         "Average of per-entry wait times across all answered entries in the selected "
         "window."),
        ("Average AHT",
         "Average of per-entry handle times (billsec + hold_accum_seconds, max leg) "
         "across answered entries."),
        ("Average MOS",
         "AVG(rtp_audio_in_mos) WHERE rtp_audio_in_mos > 0, across all CDR records "
         "in the window (not deduplicated by queue entry)."),
        ("Service Level %",
         "% of answered entries with wait ≤ 30 s. Threshold fixed at 30 s on the "
         "executive overview page."),
        ("Total Talk Time", "SUM(billsec) across all CDR records in the window."),
    ]
    add_definition_table(doc, rows)

    add_para(doc, "Trend charts on this page use SQL-level aggregations (not Python-level entry deduplication):", italic=True)
    add_bullet(doc, "Offered trend: COUNT(DISTINCT concat(caller_id_number, cc_queue_joined_epoch)) grouped by local date.")
    add_bullet(doc, "Answered trend: COUNT(DISTINCT same key WHERE cc_queue_answered_epoch IS NOT NULL) grouped by local date.")
    add_bullet(doc, "Service level trend: daily ratio of entries answered within threshold vs all answered entries.")
    add_bullet(doc, "ASA trend: AVG(cc_queue_answered_epoch − cc_queue_joined_epoch) grouped by local date.")
    add_bullet(doc, "AHT trend: AVG(billsec) WHERE cc_queue_answered_epoch IS NOT NULL grouped by local date.")

    # ======================================================================
    # SECTION 6 — CACHING
    # ======================================================================
    set_heading(doc, "6. Caching", 1)

    add_para(
        doc,
        "Several API responses are cached in Redis to reduce database load. Cache "
        "keys include all filter parameters so different filter combinations never "
        "share a cache entry.",
    )
    rows = [
        ("Executive Overview", "60 seconds TTL (configurable via EXEC_OVERVIEW_CACHE_TTL env var). Cache key bucketed by minute so near-simultaneous users share the entry."),
        ("Queue/Agent metadata", "300 seconds TTL. Cached responses for /queues-visible and /agents-visible endpoints."),
        ("Queue Performance", "Not cached — computed fresh per request."),
        ("Agent Leaderboard", "Not cached — computed fresh per request."),
    ]
    add_definition_table(doc, rows)

    # ======================================================================
    # SECTION 7 — GLOSSARY
    # ======================================================================
    set_heading(doc, "7. Glossary", 1)

    glossary = [
        ("AHT", "Average Handle Time — talk time + hold time per answered call."),
        ("ASA", "Average Speed of Answer — wait time from queue entry to agent pick-up."),
        ("Billsec", "Billable seconds — duration of the live connected portion of a call."),
        ("cc_cause", "FusionPBX call-center cause code (e.g. TIMEOUT, CANCEL)."),
        ("cc_queue_answered_epoch", "Unix epoch set by FusionPBX when a queue agent bridged the call."),
        ("cc_queue_joined_epoch", "Unix epoch set when the caller entered the queue."),
        ("CDR", "Call Detail Record — one row in the database per call leg."),
        ("Deflect", "Routing action that sends the call away without a live agent (voicemail, IVR, external transfer)."),
        ("direction", "CDR field: 'inbound' / 'outbound' / 'local'."),
        ("hold_accum_seconds", "Accumulated hold time in seconds for a call leg."),
        ("MOS", "Mean Opinion Score — voice quality score 0–5 (ITU-T P.800)."),
        ("NORMAL_CLEARING", "FusionPBX hangup cause indicating the call ended after a normal connected session."),
        ("Queue entry", "One unique caller interaction with a queue, identified by (caller_id_number, cc_queue_joined_epoch)."),
        ("Service Level", "% of answered calls answered within the configured wait threshold (default 30 s)."),
        ("Strict mode", "Queue answered classification that requires cc_queue_answered_epoch to be set."),
        ("Standard mode", "Queue answered classification that also accepts answer_epoch + NORMAL_CLEARING as evidence of an answer."),
        ("xml_cdr_uuid", "FusionPBX unique identifier for a CDR record; used as the upsert key during ETL."),
    ]
    add_definition_table(doc, glossary)

    return doc


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    output_path = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "docs",
        "CALCULATIONS.docx",
    )
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = build_document()
    doc.save(output_path)
    print(f"Saved: {output_path}")
